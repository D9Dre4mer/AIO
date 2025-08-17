# =========================================
# File Handler for Academic Papers
# Support for PDF, DOCX, and TXT file formats
# =========================================

import io
import os
from pathlib import Path
from typing import Optional, Dict, Any, List
import PyPDF2
from docx import Document


class FileHandler:
    """
    Handles file operations for academic papers.
    Supports PDF, DOCX, and TXT formats with error handling.
    """
    
    def __init__(self):
        """Initialize file handler."""
        self.supported_formats = {'.pdf', '.docx', '.txt', '.doc'}
        self.max_file_size = 50 * 1024 * 1024  # 50MB
    
    def extract_text(self, file) -> str:
        """
        Extract text from uploaded file.
        
        Args:
            file: Streamlit uploaded file object
            
        Returns:
            Extracted text content
        """
        
        if not file:
            raise ValueError("No file provided")
        
        # Get file extension
        file_extension = self._get_file_extension(file.name)
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Check file size
        if file.size > self.max_file_size:
            raise ValueError(f"File too large: {file.size / 1024 / 1024:.1f}MB")
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file)
            elif file_extension in {'.docx', '.doc'}:
                return self._extract_from_docx(file)
            elif file_extension == '.txt':
                return self._extract_from_txt(file)
            else:
                raise ValueError(f"Unsupported format: {file_extension}")
                
        except Exception as e:
            raise RuntimeError(f"Error extracting text: {str(e)}")
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        return Path(filename).suffix.lower()
    
    def _extract_from_pdf(self, file) -> str:
        """Extract text from PDF file."""
        
        try:
            # Read PDF file
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            
            # Extract text from all pages
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            # Reset file pointer
            file.seek(0)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {str(e)}")
    
    def _extract_from_docx(self, file) -> str:
        """Extract text from DOCX file."""
        
        try:
            # Read DOCX file
            doc = Document(io.BytesIO(file.read()))
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            # Reset file pointer
            file.seek(0)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise RuntimeError(f"DOCX extraction failed: {str(e)}")
    
    def _extract_from_txt(self, file) -> str:
        """Extract text from TXT file."""
        
        try:
            # Read text file
            content = file.read().decode('utf-8')
            
            # Reset file pointer
            file.seek(0)
            
            return content
            
        except UnicodeDecodeError:
            # Try different encodings
            file.seek(0)
            try:
                content = file.read().decode('latin-1')
                file.seek(0)
                return content
            except Exception as e:
                raise RuntimeError(f"Text file decoding failed: {str(e)}")
        except Exception as e:
            raise RuntimeError(f"Text file reading failed: {str(e)}")
    
    def validate_file(self, file) -> Dict[str, Any]:
        """
        Validate uploaded file.
        
        Args:
            file: Streamlit uploaded file object
            
        Returns:
            Dictionary with validation results
        """
        
        if not file:
            return {'valid': False, 'error': 'No file provided'}
        
        validation_result = {
            'valid': True,
            'filename': file.name,
            'size': file.size,
            'format': self._get_file_extension(file.name),
            'warnings': []
        }
        
        # Check file format
        if validation_result['format'] not in self.supported_formats:
            validation_result['valid'] = False
            validation_result['error'] = f"Unsupported format: {validation_result['format']}"
            return validation_result
        
        # Check file size
        if file.size > self.max_file_size:
            validation_result['valid'] = False
            validation_result['error'] = f"File too large: {file.size / 1024 / 1024:.1f}MB"
            return validation_result
        
        # Check for empty file
        if file.size == 0:
            validation_result['valid'] = False
            validation_result['error'] = 'File is empty'
            return validation_result
        
        # Add warnings for large files
        if file.size > 10 * 1024 * 1024:  # 10MB
            validation_result['warnings'].append('Large file may take longer to process')
        
        return validation_result
    
    def get_file_info(self, file) -> Dict[str, Any]:
        """
        Get detailed information about uploaded file.
        
        Args:
            file: Streamlit uploaded file object
            
        Returns:
            Dictionary with file information
        """
        
        if not file:
            return {}
        
        file_info = {
            'name': file.name,
            'size_bytes': file.size,
            'size_mb': file.size / 1024 / 1024,
            'format': self._get_file_extension(file.name),
            'type': self._get_file_type(file.name)
        }
        
        return file_info
    
    def _get_file_type(self, filename: str) -> str:
        """Get human-readable file type."""
        
        extension = self._get_file_extension(filename)
        
        type_mapping = {
            '.pdf': 'Portable Document Format',
            '.docx': 'Microsoft Word Document',
            '.doc': 'Microsoft Word Document (Legacy)',
            '.txt': 'Plain Text File'
        }
        
        return type_mapping.get(extension, 'Unknown')
    
    def batch_process_files(self, files: List) -> List[Dict[str, Any]]:
        """
        Process multiple files in batch.
        
        Args:
            files: List of uploaded files
            
        Returns:
            List of processing results
        """
        
        results = []
        
        for file in files:
            try:
                # Validate file
                validation = self.validate_file(file)
                
                if not validation['valid']:
                    results.append({
                        'filename': file.name,
                        'success': False,
                        'error': validation['error']
                    })
                    continue
                
                # Extract text
                text_content = self.extract_text(file)
                
                results.append({
                    'filename': file.name,
                    'success': True,
                    'text_length': len(text_content),
                    'word_count': len(text_content.split()),
                    'format': validation['format']
                })
                
            except Exception as e:
                results.append({
                    'filename': file.name,
                    'success': False,
                    'error': str(e)
                })
        
        return results
    
    def save_extracted_text(self, text: str, filename: str, 
                           output_dir: str = "extracted_texts") -> str:
        """
        Save extracted text to file.
        
        Args:
            text: Extracted text content
            filename: Original filename
            output_dir: Output directory
            
        Returns:
            Path to saved file
        """
        
        # Create output directory
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # Generate output filename
        base_name = Path(filename).stem
        output_file = output_path / f"{base_name}_extracted.txt"
        
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            return str(output_file)
            
        except Exception as e:
            raise RuntimeError(f"Failed to save extracted text: {str(e)}")
    
    def create_summary_report(self, processing_results: List[Dict[str, Any]]) -> str:
        """
        Create a summary report of file processing.
        
        Args:
            processing_results: List of processing results
            
        Returns:
            Formatted summary report
        """
        
        if not processing_results:
            return "No files processed."
        
        total_files = len(processing_results)
        successful_files = sum(1 for r in processing_results if r['success'])
        failed_files = total_files - successful_files
        
        # Calculate total text extracted
        total_text_length = sum(r.get('text_length', 0) 
                               for r in processing_results if r['success'])
        total_words = sum(r.get('word_count', 0) 
                         for r in processing_results if r['success'])
        
        # Format summary
        summary = f"""
File Processing Summary
======================

Total Files: {total_files}
Successful: {successful_files}
Failed: {failed_files}

Text Extracted:
- Total Characters: {total_text_length:,}
- Total Words: {total_words:,}

Files by Format:
"""
        
        # Count files by format
        format_counts = {}
        for result in processing_results:
            if result['success']:
                format_type = result.get('format', 'unknown')
                format_counts[format_type] = format_counts.get(format_type, 0) + 1
        
        for format_type, count in format_counts.items():
            summary += f"- {format_type.upper()}: {count}\n"
        
        # Add failed files details
        if failed_files > 0:
            summary += f"\nFailed Files:\n"
            for result in processing_results:
                if not result['success']:
                    summary += f"- {result['filename']}: {result['error']}\n"
        
        return summary.strip()
